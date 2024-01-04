from docx import Document
import pandas as pd
from datetime import datetime
from docx.shared import Inches, Pt, RGBColor  
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import _Cell
from docx.enum.table import WD_ROW_HEIGHT_RULE
import os
import sys
import folium
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import tempfile

# Función para agregar una cabecera al documento
def agregar_cabecera(doc):
    cabecera = doc.sections[0].header  # Obtener la cabecera de la primera sección del documento
    cabecera.paragraphs[0].clear()  # Limpiar el contenido existente en la cabecera

    # Agregar el texto de la cabecera
    parrafo_cabecera = cabecera.paragraphs[0]
    texto_cabecera = "CONTRACTACIÓ PER LA PRESENTACIÓ DEL SERVEI D'UNITAT D'INTERVENCIÓ RÀPIDA, PEL MANTENIMENT DE LA VIA PÚBLICA I DELS EDIFICIS MUNICIPALS"
    parrafo_cabecera.text = texto_cabecera.upper()  # Texto de la cabecera en mayúsculas
    parrafo_cabecera.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinear el texto al centro
    parrafo_cabecera.runs[0].font.bold = True  # Aplicar negrita al texto
    parrafo_cabecera.runs[0].font.size = Pt(10)  # Tamaño de fuente
    parrafo_cabecera.runs[0].font.color.rgb = RGBColor(0x00, 0x70, 0xC0)  # Color azul ##0070C0

def agregar_pie_de_pagina(doc):
    pie_de_pagina = doc.sections[0].footer  # Obtener el pie de página de la primera sección del documento
    pie_de_pagina.paragraphs[0].clear()  # Limpiar el contenido existente en el pie de página

    # Crear un nuevo párrafo en el pie de página
    parrafo = pie_de_pagina.add_paragraph()

    # Agregar el texto "SOBRE A. CRITERIS AVALUABLES A JUDICI DE VALOR" al párrafo
    run = parrafo.add_run("SOBRE B. CRITERIS AVALUABLES A JUDICI DE VALOR")
    run.font.bold = True  # Aplicar negrita al texto
    run.font.size = Pt(9)  # Tamaño de fuente
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)  # Color azul #0088FF
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinear el texto al centro

def set_background_color(cell, color):
    # Obtener el elemento "tcPr" que representa las propiedades de la celda
    tc_pr = cell._element.get_or_add_tcPr()

    # Crear un elemento "shd" para establecer el color de fondo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def set_cell_border(cell: _Cell, border_color: str):
    # Establecer el color del borde de una celda
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Crear elemento para bordes de celda
    tcBorders = parse_xml(r'<w:tcBorders {}>'.format(nsdecls('w')) +
                          '<w:top w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(border_color) +
                          '<w:left w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(border_color) +
                          '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(border_color) +
                          '<w:right w:val="single" w:sz="4" w:space="0" w:color="{}"/>'.format(border_color) +
                          '</w:tcBorders>')
    tcPr.append(tcBorders)


def set_font_style(element, font_name):
    element.rPr.rFonts.set(qn('w:ascii'), font_name)
    element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    element.rPr.rFonts.set(qn('w:cs'), font_name)

def apply_font_format(cell, bold=False, font_color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=12):
    paragraph = cell.paragraphs[0]
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    if bold:
        run.bold = bold
    if font_color:
        run.font.color.rgb = font_color
    paragraph.alignment = alignment

def leer_datos_desde_excel(ruta_excel, ruta_excel_calles):
    # Leer el archivo Excel y cargarlo en un DataFrame de pandas
    df = pd.read_excel(ruta_excel)

    df_calles = pd.read_excel(ruta_excel_calles)
    df['barri'] = ''

    for index, row in df.iterrows():
        lloc_thoroughfare = row['lloc_thoroughfare'].lower() if pd.notna(row['lloc_thoroughfare']) else ''
        for _, calle_row in df_calles.iterrows():
            nom_via = calle_row['NOM_VIA'].lower() if pd.notna(calle_row['NOM_VIA']) else ''
            if nom_via in lloc_thoroughfare:
                df.at[index, 'barri'] = calle_row['BARRI']
                break

    # Lista para almacenar los datos de cada fila
    datos_filas = []

    # Leer los datos de cada fila y agregarlos a la lista
    for _, fila in df.iterrows():
        # Recogiendo los datos necesarios de la fila
        titulo = fila["_title"]
        lloc = str(fila["lloc"])
        lloc_thoroughfare = fila["lloc_thoroughfare"] 
        fecha = fila["data"]
        imatges = fila["imatges"]
        latitud = fila["_latitude"]
        longitud = fila["_longitude"]
        barri = fila["barri"] 
        edifici = fila["edifici"] 
        sala = fila["sala"] 
        numero_de_planta = fila["numero_de_planta"]

        #Num rubatec 
        
        num_incidencia = fila["num_incidencia"]    

        # Recogiendo los tipos de operaciones y las interferencias, solo si existen (no son NaN)
        desperfectos = []
        amidament = []
        unitats = []
        interferencias = []
        propuestas = []
        for i in range(1, 4):
            if pd.notna(fila[f"{i}_tipus_de_desperfecte"]):
                desperfectos.append(fila[f"{i}_tipus_de_desperfecte"])
            if pd.notna(fila[f"{i}_amidament"]):
                amidament.append(fila[f"{i}_amidament"])
            if pd.notna(fila[f"{i}_unitats"]):
                unitats.append(fila[f"{i}_unitats"])           
            if pd.notna(fila[f"{i}_tipus_operacio"]):
                propuestas.append(fila[f"{i}_tipus_operacio"])            
            if pd.notna(fila[f"{i}_interferencia"]):
                interferencias.append(fila[f"{i}_interferencia"])
        # Agregar la fila de datos a la lista
        datos_filas.append((
            titulo, lloc, lloc_thoroughfare, fecha, imatges, latitud, longitud, barri,
            edifici, sala, numero_de_planta,desperfectos, amidament,unitats, propuestas , interferencias, num_incidencia
        ))
        
    return datos_filas

def agregar_imagen_en_celda(celda, ruta_imagen):
    # Crear un párrafo en la celda
    paragraph = celda.paragraphs[0]
    run = paragraph.add_run()
    # Añadir la imagen al párrafo de la celda
    run.add_picture(ruta_imagen, width=Inches(3))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generar_mapa_folium(latitud, longitud):
    mapa = folium.Map(location=[latitud, longitud], max_zoom=19, zoom_start=19)
    
    # Agregar un marcador en la ubicación indicada por las coordenadas de latitud y longitud
    folium.Marker([latitud, longitud]).add_to(mapa)

    # Guardar el mapa en un archivo HTML temporal
    mapa_html = tempfile.NamedTemporaryFile(suffix=".html", delete=False)
    mapa.save(mapa_html.name)
    mapa_html.close()

    return mapa_html.name

def crear_driver_web():
    # Determinar la ruta del directorio actual del script o ejecutable
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))

    # Agregar la ruta del directorio al PATH
    os.environ['PATH'] += os.pathsep + base_path

    # Configurar opciones para Chrome
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')

    # Crear y retornar el driver de Chrome sin especificar la ruta
    return webdriver.Chrome(options=options)

def crear_bloque_desperfecto(tabla, i, desperfecto, amidament,unitats, interferencia, propuesta):
    # Asumiendo que 'tabla' es un objeto de tabla de python-docx
    titles = ["Desperfecte", "Amidament", "Interferència/afecció amb altres usos públics", "Proposta d'activitat"]
    values = [desperfecto, f"{amidament} {unitats}", interferencia, propuesta]

    for title, value in zip(titles, values):
        row_cells = tabla.add_row().cells
        row_cells[0].text = f"{title} {i+1}" if title == "Desperfecte" else title
        set_background_color(row_cells[0], '548DD4')
        apply_font_format(row_cells[0], bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        row_cells[1].text = str(value)
        set_background_color(row_cells[1], 'BFBFBF')
        apply_font_format(row_cells[1], alignment=WD_ALIGN_PARAGRAPH.CENTER)  # Formato por defecto en negro BFBFBF

def add_building_info(tabla, edifici, sala, numero_de_planta):
    if pd.isna(edifici):
        return

    info_titles = ["EDIFICI", "SALA", "PLANTA"]
    info_values = [edifici, sala, numero_de_planta]

    for title, value in zip(info_titles, info_values):
        row = tabla.add_row().cells
        row[0].text = title
        set_background_color(row[0], '548DD4')
        apply_font_format(row[0], bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        row[1].text = str(value)
        set_background_color(row[1], 'BFBFBF')
        apply_font_format(row[1], alignment=WD_ALIGN_PARAGRAPH.CENTER)   

def agregar_imagen_mapa_barrio(doc, barri, carpeta_mapas_barrios):
    for archivo in os.listdir(carpeta_mapas_barrios):
        if archivo.lower().endswith('.jpg') and barri.lower() in archivo.lower():
            ruta_imagen_barrio = os.path.join(carpeta_mapas_barrios, archivo)
            if os.path.exists(ruta_imagen_barrio):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(ruta_imagen_barrio, width=Inches(3.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break  # Salir del bucle una vez que se encuentra y agrega la imagen correspondiente   

def obtener_ruta_imagen_barrio(barri, carpeta_mapas_barrios):
    for archivo in os.listdir(carpeta_mapas_barrios):
        if archivo.lower().endswith('.jpg') and barri.lower() in archivo.lower():
            return os.path.join(carpeta_mapas_barrios, archivo)
    return None            


def crear_tablas_informes(datos_filas,carpeta_imagenes,carpeta_mapas_barrios):
    # Crear el objeto Document
    doc = Document()

    # Crear una tabla para cada fila de datos
    for datos in datos_filas:
        
        titulo, lloc, lloc_thoroughfare, fecha, imatges, latitud, longitud, barri, edifici, sala, numero_de_planta,desperfectos,amidaments,unitats, propuestas , interferencias, num_incidencia = datos

        # Dividir el título en elementos
        elementos = titulo.split(", ")

        # Crear la tabla con 5 filas y 2 columnas
        tabla = doc.add_table(rows=9, cols=2)
        tabla.style = 'Table Grid'
        
        # Configuración de las celdas de la tabla
        for row in tabla.rows:
            for cell in row.cells:
                cell.width = Inches(3.5)
        
        # Primera fila - "INCIDÈNCIA"
        celda_encabezado = tabla.cell(0, 0)
        celda_encabezado.merge(tabla.cell(0, 1))
        celda_encabezado.text = "INCIDÈNCIA Nº "+ str(num_incidencia) 
        set_background_color(celda_encabezado, '002060')
        apply_font_format(celda_encabezado, bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=18)

        # Fila fantasma

        celda_fantasma = tabla.cell(1,0)
        celda_fantasma.merge(tabla.cell(1,1))

        #TITULO DE LA SEGUNDA FILA - DESCRIPCIÓ DE L'INCIDENCIA

        celda_titulo_descripcio = tabla.cell(2,0)
        celda_titulo_descripcio.merge(tabla.cell(2,1))
        celda_titulo_descripcio.text = "DESCRIPCIÓ INCIDÈNCIA"
        set_background_color(celda_titulo_descripcio, '002060')
        apply_font_format(celda_titulo_descripcio, bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=15)

        #  Descripción de la incidencia
        celda_descripcion = tabla.cell(3, 0)
        celda_descripcion.merge(tabla.cell(3, 1))
        descripcion_incidencia = ", ".join([f"{elem} {desp}" for elem, desp in zip(elementos, desperfectos)]) + f" a {lloc_thoroughfare}"
        celda_descripcion.text = str(descripcion_incidencia)
        apply_font_format(celda_descripcion)  # Fuente por defecto en negro

        #TITULO DE LOS DATOS - DADES

        celda_titulo_dades = tabla.cell(4,0)
        celda_titulo_dades.merge(tabla.cell(4,1))
        celda_titulo_dades.text = "DADES"
        set_background_color(celda_titulo_dades, '002060')
        apply_font_format(celda_titulo_dades, bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=15)

        #  "DATA"
        celda_data = tabla.cell(5, 0)
        celda_data.text = "DATA DETECCIÓ"
        set_background_color(celda_data, '548DD4')
        apply_font_format(celda_data, bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        tabla.cell(5, 1).text = str(fecha)
        set_background_color(tabla.cell(5, 1), 'BFBFBF')
        apply_font_format(tabla.cell(5, 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)  # Fuente por defecto en negro

        # "BARRI"
        celda_barri = tabla.cell(6, 0)
        celda_barri.text = "BARRI"
        set_background_color(celda_barri, '548DD4')
        apply_font_format(celda_barri, bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        tabla.cell(6, 1).text = str(barri)
        set_background_color(tabla.cell(6, 1), 'BFBFBF')
        apply_font_format(tabla.cell(6, 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)  # Fuente por defecto en negro

        # "LOCALITZACIÓ"
        celda_localitzacio = tabla.cell(7, 0)
        celda_localitzacio.text = "LOCALITZACIÓ"
        set_background_color(celda_localitzacio, '548DD4')
        apply_font_format(celda_localitzacio, bold=True, font_color=RGBColor(255, 255, 255),alignment=WD_ALIGN_PARAGRAPH.CENTER)
        tabla.cell(7, 1).text = str(lloc_thoroughfare)
        set_background_color(tabla.cell(7, 1), 'BFBFBF')
        apply_font_format(tabla.cell(7, 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)  # Fuente por defecto en negro

        # ELEMENT AFECTAT
        celda_elem = tabla.cell(8, 0)
        celda_elem.text = "ELEMENT AFECTAT"
        set_background_color(celda_elem, '548DD4')
        apply_font_format(celda_elem, bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        tabla.cell(8, 1).text = str(titulo)
        set_background_color(tabla.cell(8, 1), 'BFBFBF')
        apply_font_format(tabla.cell(8, 1), alignment=WD_ALIGN_PARAGRAPH.CENTER)  # Fuente por defecto en negro

        
        add_building_info(tabla, edifici, sala, numero_de_planta)

        for i, (desperfecto, amdt, unit, interf, prop) in enumerate(zip(desperfectos, amidaments, unitats, interferencias, propuestas)):
            crear_bloque_desperfecto(tabla, i, desperfecto, amdt, unit, interf, prop)

        #INCIDÈNCIA GRÀFICA

        celdas_titulo_incidencia = tabla.add_row().cells
        celda_titulo_incidencia = celdas_titulo_incidencia[0].merge(celdas_titulo_incidencia[1])
        celda_titulo_incidencia.text = "INCIDÈNCIA GRÀFICA"
        set_background_color(celda_titulo_incidencia, '002060')
        apply_font_format(celda_titulo_incidencia, bold=True, font_size=15, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            

        # Agregar imágenes dentro de la tabla
        identificadores_imagenes = imatges.split(",")[:2]  # Tomar los dos primeros identificadores
        if len(identificadores_imagenes) > 0:
            row_imagenes = tabla.add_row().cells
            for i, identificador in enumerate(identificadores_imagenes):
                ruta_imagen = os.path.join(carpeta_imagenes, f"{identificador}.jpg")
                if os.path.exists(ruta_imagen):
                    run = row_imagenes[i].paragraphs[0].add_run()
                    run.add_picture(ruta_imagen, height=Inches(3.0))
                    row_imagenes[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        #--------------------------------------------------------------

        # Crear la tabla 
        tabla2 = doc.add_table(rows=7, cols=2)
        tabla.style = 'Table Grid'

        for row in tabla2.rows:
            for cell in row.cells:
                set_cell_border(cell, "FFFFFF")

        coordenades = tabla2.cell(0,0).merge(tabla2.cell(0,1))
        coordenades.text = "COORDENADES"
        set_background_color(coordenades, '002060')
        apply_font_format(coordenades, bold=True, font_color=RGBColor(255, 255, 255), font_size=15,alignment=WD_ALIGN_PARAGRAPH.CENTER)


        tabla2.cell(1,0).text = "LATITUD" 
        set_background_color(tabla2.cell(1,0), '548DD4')
        apply_font_format(tabla2.cell(1,0), bold=True, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        tabla2.cell(1,1).text = str(latitud) 
        set_background_color(tabla2.cell(1,1), 'BFBFBF')
        apply_font_format(tabla2.cell(1,1), bold=True,  alignment=WD_ALIGN_PARAGRAPH.CENTER)

        tabla2.cell(2,0).text = "LONGITUD"
        set_background_color(tabla2.cell(2,0), '548DD4')
        apply_font_format(tabla2.cell(2,0), bold=True,font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)

        tabla2.cell(2,1).text = str(longitud)
        set_background_color(tabla2.cell(2,1), 'BFBFBF')
        apply_font_format(tabla2.cell(2,1), bold=True,  alignment=WD_ALIGN_PARAGRAPH.CENTER)

        loc_geo = tabla2.cell(3,0).merge(tabla2.cell(3,1))
        loc_geo.text = "LOCALITZACIÓ GRÀFICA"
        set_background_color(loc_geo, '002060')
        apply_font_format(loc_geo, bold=True,font_size=13, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)

        tabla2.cell(4,0).merge(tabla2.cell(4,1))

        # AQUI DEBERIA PONER LA FOTO DE MAPA BARRIO
        ruta_imagen_barrio = obtener_ruta_imagen_barrio(barri, carpeta_mapas_barrios)
        if ruta_imagen_barrio:
            agregar_imagen_en_celda(tabla2.cell(4, 0).merge(tabla2.cell(4, 1)), ruta_imagen_barrio)

        geolocalitzacio = tabla2.cell(5,0).merge(tabla2.cell(5,1))
        geolocalitzacio.text = "GEOLOCALITZACIÓ"
        set_background_color(geolocalitzacio, '002060')
        apply_font_format(geolocalitzacio, bold=True, font_size=13, font_color=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.CENTER)

        #AQUI DEBERIA PONER LA FOTO DE FOLIUM
 
        # Generar y agregar el mapa
        print("Generando mapa con Folium")
        mapa_html_file = generar_mapa_folium(latitud, longitud)
        print("Mapa generado: ", mapa_html_file)
        options = Options()
        options.add_argument('--headless')
        driver = crear_driver_web()
        driver.get("file:///" + mapa_html_file)
        driver.set_window_size(800, 600)
        ruta_imagen_geolocalizacion = tempfile.mktemp(suffix=".png")
        driver.save_screenshot(ruta_imagen_geolocalizacion)
        driver.quit()
        os.remove(mapa_html_file)

        agregar_imagen_en_celda(tabla2.cell(6,0).merge(tabla2.cell(6,1)), ruta_imagen_geolocalizacion) 

        os.remove(ruta_imagen_geolocalizacion)

        # Agregar un salto de página después de cada tabla
        doc.add_page_break()

    # Agregar cabecera y pie de página
    agregar_cabecera(doc)
    agregar_pie_de_pagina(doc)

    # Guardar el documento en un archivo
    doc.save("informes_word.docx")

import tkinter as tk
from tkinter import filedialog

def generar_informes(ruta_excel, carpeta_imagenes, ruta_excel_calles, carpeta_mapas_barrios):
    print("Inicio de generar_informes")
    datos_filas = leer_datos_desde_excel(ruta_excel, ruta_excel_calles)
    crear_tablas_informes(datos_filas, carpeta_imagenes, carpeta_mapas_barrios)
    print("Fin de generar_informes")


def main():
    def browse_file():
        filename = filedialog.askopenfilename()
        entry_excel_path.delete(0, tk.END)
        entry_excel_path.insert(0, filename)

    def browse_folder():
        foldername = filedialog.askdirectory()
        entry_photos_path.delete(0, tk.END)
        entry_photos_path.insert(0, foldername)

    def browse_streets_excel():
        filename = filedialog.askopenfilename()
        entry_streets_excel_path.delete(0, tk.END)
        entry_streets_excel_path.insert(0, filename)

    def browse_barrios_map_folder():
        foldername = filedialog.askdirectory()
        entry_barrios_map_path.delete(0, tk.END)
        entry_barrios_map_path.insert(0, foldername)

    def execute_script():
        ruta_excel = entry_excel_path.get()
        carpeta_imagenes = entry_photos_path.get()
        ruta_excel_calles = entry_streets_excel_path.get()
        carpeta_mapas_barrios = entry_barrios_map_path.get()
        # Aquí llamas a tu función principal del script con las rutas y carpetas obtenidas
        generar_informes(ruta_excel, carpeta_imagenes, ruta_excel_calles, carpeta_mapas_barrios)

    root = tk.Tk()
    root.title("Generador de Informes")

    # Configuración de la GUI para seleccionar el archivo Excel principal
    tk.Label(root, text="Ruta al archivo Excel principal:").pack()
    entry_excel_path = tk.Entry(root, width=50)
    entry_excel_path.pack()
    tk.Button(root, text="Buscar", command=browse_file).pack()

    # Configuración de la GUI para seleccionar la carpeta de fotos
    tk.Label(root, text="Ruta a la carpeta de fotos:").pack()
    entry_photos_path = tk.Entry(root, width=50)
    entry_photos_path.pack()
    tk.Button(root, text="Buscar", command=browse_folder).pack()

    # Configuración de la GUI para seleccionar el archivo Excel de calles y barrios
    tk.Label(root, text="Ruta al archivo Excel de calles y barrios:").pack()
    entry_streets_excel_path = tk.Entry(root, width=50)
    entry_streets_excel_path.pack()
    tk.Button(root, text="Buscar", command=browse_streets_excel).pack()

    # Configuración de la GUI para seleccionar la carpeta de mapas de barrios
    tk.Label(root, text="Ruta a la carpeta de mapas de barrios:").pack()
    entry_barrios_map_path = tk.Entry(root, width=50)
    entry_barrios_map_path.pack()
    tk.Button(root, text="Buscar", command=browse_barrios_map_folder).pack()

    # Botón para ejecutar el script
    tk.Button(root, text="Generar Informes", command=execute_script).pack()

    root.mainloop()

if __name__ == "__main__":
    main()


