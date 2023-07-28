from docx import Document
import pandas as pd
from datetime import datetime
from docx.shared import Inches, Pt, RGBColor  
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import folium
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import tempfile

# Función para agregar una cabecera al documento
def agregar_cabecera(doc):
    cabecera = doc.sections[0].header  # Obtener la cabecera de la primera sección del documento
    cabecera.paragraphs[0].clear()  # Limpiar el contenido existente en la cabecera

    # Agregar el texto de la cabecera
    parrafo_cabecera = cabecera.paragraphs[0]
    texto_cabecera = "CONTRACTACIÓ PER LA PRESENTACIÓ DEL SERVEI D'UNITAT D'INTERVENCIÓ RÀPIDA, PEL MANTENIMENT DE LA VIA PÚBLICA I DELS EDIFICIS MUNICIPALS"
    parrafo_cabecera.text = texto_cabecera.upper()  # Texto de la cabecera en mayúsculas
    parrafo_cabecera.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinear el texto al centro
    parrafo_cabecera.runs[0].font.bold = True  # Aplicar negrita al texto
    parrafo_cabecera.runs[0].font.size = Pt(10)  # Tamaño de fuente
    parrafo_cabecera.runs[0].font.color.rgb = RGBColor(0x00, 0x70, 0xC0)  # Color azul ##0070C0

def agregar_pie_de_pagina(doc):
    pie_de_pagina = doc.sections[0].footer  # Obtener el pie de página de la primera sección del documento
    pie_de_pagina.paragraphs[0].clear()  # Limpiar el contenido existente en el pie de página

    # Crear un nuevo párrafo en el pie de página
    parrafo = pie_de_pagina.add_paragraph()

    # Agregar el texto "SOBRE A. CRITERIS AVALUABLES A JUDICI DE VALOR" al párrafo
    run = parrafo.add_run("SOBRE A. CRITERIS AVALUABLES A JUDICI DE VALOR")
    run.font.bold = True  # Aplicar negrita al texto
    run.font.size = Pt(9)  # Tamaño de fuente
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)  # Color azul #0088FF
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinear el texto al centro

def set_background_color(cell, color):
    # Obtener el elemento "tcPr" que representa las propiedades de la celda
    tc_pr = cell._element.get_or_add_tcPr()

    # Crear un elemento "shd" para establecer el color de fondo
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

def set_font_style(element, font_name):
    element.rPr.rFonts.set(qn('w:ascii'), font_name)
    element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    element.rPr.rFonts.set(qn('w:cs'), font_name)

def apply_font_format(cell, bold=False, font_color=None):
    # Obtener el párrafo de la celda
    paragraph = cell.paragraphs[0]

    # Obtener el elemento "r" que representa el texto del párrafo
    run = paragraph.runs[0]

    # Aplicar formato de fuente

    run.font.size = Pt(11)

    if bold==True:  
        run.bold = bold

    # Configurar fuente Arial
    run.font.name = "Arial"
    if font_color:
        run.font.color.rgb = font_color

def leer_datos_desde_excel(ruta_excel):
    # Leer el archivo Excel y cargarlo en un DataFrame de pandas
    df = pd.read_excel(ruta_excel)

    # Lista para almacenar los datos de cada fila
    datos_filas = []

    # Leer los datos de cada fila y agregarlos a la lista
    for _, fila in df.iterrows():
        titulo = fila["_title"]
        valores_columnas = [str(fila[col]) for col in df.columns if str(col).lower()[0] in ["v", "e", "c"] and pd.notna(fila[col])]
        #lloc2 = str(fila["lloc"])  # Variable separada para almacenar el valor de la columna "lloc" del Excel
        datos_filas.append((titulo, valores_columnas, str(fila["lloc_thoroughfare"]), str(fila["lloc"]),
                            fila["1_amidament"], fila["1_unitats"],
                            fila["2_amidament"], fila["2_unitats"],
                            fila["3_amidament"], fila["3_unitats"],
                            fila["data"],fila["imatges"],
                            fila["_latitude"],fila["_longitude"]))
        
    #print(datos_filas)    

    return datos_filas #,lloc2

def combinar_strings(lista1, lista2):
    # Combinar las listas en el formato deseado
    combinacion = []
    for i, item in enumerate(lista1):
        valor = lista2[i] if i < len(lista2) else ""
        combinacion.append(f"{item} {valor}")
    return ", ".join(combinacion)

def generar_mapa_folium(latitud, longitud):
    mapa = folium.Map(location=[latitud, longitud], max_zoom=19, zoom_start=19)
    
    # Agregar un marcador en la ubicación indicada por las coordenadas de latitud y longitud
    folium.Marker([latitud, longitud]).add_to(mapa)

    # Guardar el mapa en un archivo HTML temporal
    mapa_html = tempfile.NamedTemporaryFile(suffix=".html", delete=False)
    mapa.save(mapa_html.name)
    mapa_html.close()

    return mapa_html.name

def crear_tablas_informes(datos_filas):
    # Crear el objeto Document
    doc = Document()

    # Datos para la tabla de la columna izquierda
    valores_columna_izquierda = [
        "Incidència",
        "Primer desperfecte",
        "Segon desperfecte",
        "Tercer desperfecte",
        "Data",
        "Lloc"
    ]

    carpeta_imagenes = r"C:\Users\itacl\stjust\data\photos"

    # Crear una tabla para cada fila de datos
    for datos_fila in datos_filas:
        titulo, valores_columnas, lloc, lloc2, amidament1, unitats1, amidament2, unitats2, amidament3, unitats3, fecha, imatges, latitud, longitud = datos_fila

        # Crear la tabla
        filas = 7  # Se ajusta a 7 para tener 6 filas de datos y la fila de encabezados
        columnas = 2
        tabla = doc.add_table(rows=filas, cols=columnas)

        tabla.style = 'Table Grid'

        # Llenar la columna izquierda de la tabla
        for i, valor in enumerate(valores_columna_izquierda):
            celda = tabla.cell(i, 0)  # Columna izquierda está en la posición 0
            celda.text = valor

        # Combinar los valores de la columna derecha
        valores_combinados = combinar_strings(titulo.split(", "), valores_columnas)

        # Agregar el valor de la columna "lloc" al final de la celda (1, 2)
        valores_combinados += f" al {lloc}"

        # Llenar la posición (1, 1) con formato

        set_background_color(tabla.cell(0, 0), "002060")

        apply_font_format(tabla.cell(0, 0), bold=True, font_color=RGBColor(255, 255, 255))

        # Posición (1, 2) de la tabla

        celda_derecha = tabla.cell(0, 1)  
        celda_derecha.text = valores_combinados
        apply_font_format(tabla.cell(0, 1), font_color=RGBColor(0, 0, 0))

        # Llenar la posición (2, 1) con formato

        set_background_color(tabla.cell(1, 0), "002060")

        apply_font_format(tabla.cell(1, 0), bold=True, font_color=RGBColor(255, 255, 255))

        # Llenar la posición (2, 2) con los valores de "1_amidament" y "1_unitats"
        tabla.cell(1, 1).text = f"{amidament1} {unitats1}"
        apply_font_format(tabla.cell(1, 1), font_color=RGBColor(0, 0, 0))

        # Llenar la posición (3, 1) con formato

        set_background_color(tabla.cell(2, 0), "002060")

        apply_font_format(tabla.cell(2, 0), bold=True, font_color=RGBColor(255, 255, 255))

        # Llenar la posición (3, 2) con los valores de "2_amidament" y "2_unitats"
        tabla.cell(2, 1).text = f"{amidament2} {unitats2}"
        apply_font_format(tabla.cell(2, 1), font_color=RGBColor(0, 0, 0))

        # Llenar la posición (4, 1) con formato

        set_background_color(tabla.cell(3, 0), "002060")

        apply_font_format(tabla.cell(3, 0), bold=True, font_color=RGBColor(255, 255, 255))

        # Llenar la posición (4, 2) con los valores de "3_amidament" y "3_unitats"
        tabla.cell(3, 1).text = f"{amidament3} {unitats3}"
        apply_font_format(tabla.cell(3, 1), font_color=RGBColor(0, 0, 0))
        
        # Llenar la posición (5, 1) con formato

        set_background_color(tabla.cell(4, 0), "002060")

        apply_font_format(tabla.cell(4, 0), bold=True, font_color=RGBColor(255, 255, 255))

        # Llenar la posición (5, 2) con la fecha en formato dd/mm/yyyy

        fecha_datetime = datetime.strptime(fecha, "%Y-%m-%d")  # Convertir la cadena fecha a un objeto datetime
        tabla.cell(4, 1).text = fecha_datetime.strftime("%d/%m/%Y")
        apply_font_format(tabla.cell(4, 1), font_color=RGBColor(0, 0, 0))

        # Llenar la posición (6, 1) con formato

        set_background_color(tabla.cell(5, 0), "002060")

        apply_font_format(tabla.cell(5, 0), bold=True, font_color=RGBColor(255, 255, 255))

        # Llenar la posición (6, 2) con el valor de "lloc"
        tabla.cell(5, 1).text = lloc2
        apply_font_format(tabla.cell(5, 1), font_color=RGBColor(0, 0, 0))

        # Leer los identificadores de imágenes de la columna "imatges"
        identificadores_imagenes = imatges.split(",")[:2]  # Tomar los dos primeros identificadores

        # Redimensionar e insertar las imágenes en la celda (7,1) y (7,2) de la tabla
        for i, identificador in enumerate(identificadores_imagenes):
            ruta_imagen = os.path.join(carpeta_imagenes, f"{identificador}.jpg")
            if os.path.exists(ruta_imagen):
                celda_imagen = tabla.cell(6, i)  # Fila 7
                celda_imagen.paragraphs[0].add_run("\n")
                celda_imagen.paragraphs[0].alignment = 1  # Alinear imagen al centro de la celda
                celda_imagen.paragraphs[0].add_run().add_picture(ruta_imagen, height=Inches(2.0))  # Ajustar el tamaño de la imagen (2.0 pulgadas)
                celda_imagen.paragraphs[0].add_run("\n")

        # Eliminar la fila correspondiente si 2_amidament o 3_amidament es NaN
        if pd.isna(amidament3):
            tabla._tbl.remove(tabla.rows[3]._tr)       
        
        if pd.isna(amidament2):
            tabla._tbl.remove(tabla.rows[2]._tr)

        # Agregar un salto de línea después de cada tabla
        doc.add_paragraph("")  # Agregar un salto de línea    

        # Generar el mapa de Folium y obtener la ruta del archivo HTML temporal
        mapa_html_file = generar_mapa_folium(latitud, longitud)

        # Crear una captura de pantalla del mapa utilizando Selenium WebDriver y guardarla como archivo temporal
        options = Options()
        options.add_argument('--headless')  # Ejecutar el navegador en modo silencioso (sin ventana)
        driver = webdriver.Chrome(options=options)

        driver.get("file:///" + mapa_html_file)  # Abrir el archivo HTML local
        driver.set_window_size(800, 600)  # Ajustar el tamaño de la ventana para la captura de pantalla

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmpfile:
            ruta_imagen_geolocalizacion = tmpfile.name
            driver.save_screenshot(ruta_imagen_geolocalizacion)

        driver.quit()

        # Eliminar el archivo HTML temporal después de usarlo
        os.remove(mapa_html_file)

        # Agregar la imagen al documento Word después de cada tabla
        doc.add_picture(ruta_imagen_geolocalizacion, width=Inches(4.0))  # Ajustar el tamaño de la imagen (6.0 pulgadas)
        ultimo_parrafo = doc.paragraphs[-1]
        ultimo_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinear el párrafo al centro del documento

        # Agregar un salto de página después de cada tabla
        doc.add_page_break()

    # Agregar cabecera y pie de página

    agregar_cabecera(doc)
    agregar_pie_de_pagina(doc)

    # Guardar el documento en un archivo

    doc.save("informes_word.docx")

if __name__ == "__main__":
    # Ruta del archivo Excel
    ruta_excel = r"C:\Users\itacl\stjust\data\appauditoria.xlsx"

    # Leer los datos del Excel
    datos_filas = leer_datos_desde_excel(ruta_excel)

    # Crear los informes en Word con los datos obtenidos del Excel
    crear_tablas_informes(datos_filas)
